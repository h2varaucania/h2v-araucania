#!/usr/bin/env python3
"""Genera Manual_Usuario_H2V_Araucania.docx desde fuente/manual.md (python-docx).
Uso: python3 docs/manual-usuario/herramientas/build_docx.py
"""
import pathlib, datetime
from PIL import Image
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from manual_parser import parse, parse_inline, ADMON_LABEL

AQUI = pathlib.Path(__file__).resolve().parent
BASE = AQUI.parent
FIG = BASE / 'figuras'
SALIDA = BASE / 'Manual_Usuario_H2V_Araucania.docx'

AZUL = RGBColor(0x1B, 0x3A, 0x5C)
VERDE = RGBColor(0x0D, 0x73, 0x77)
GRIS = RGBColor(0x55, 0x5F, 0x6B)
ADMON_COLOR = {'consejo': 'E3F3F0', 'atencion': 'FFF1DB', 'resultado': 'E8F5E9', 'nota': 'E8F0FB', 'seguridad': 'FDE8E8'}
ADMON_BORDE = {'consejo': '0D7377', 'atencion': 'D98A00', 'resultado': '2E7D32', 'nota': '1B3A5C', 'seguridad': 'C62828'}

def set_cell_shading(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)

def set_cell_borders(cell, color='BFC5CC', sz=4, left_color=None, left_sz=None):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}'); b.set(qn('w:val'), 'single')
        if side == 'left' and left_color:
            b.set(qn('w:sz'), str(left_sz or 24)); b.set(qn('w:color'), left_color)
        else:
            b.set(qn('w:sz'), str(sz)); b.set(qn('w:color'), color)
        b.set(qn('w:space'), '0'); borders.append(b)
    tcPr.append(borders)

def add_runs(par, text, size=None, color=None, bold_all=False):
    for kind, t in parse_inline(text):
        r = par.add_run(t)
        if kind == 'bold' or bold_all: r.bold = True
        if kind == 'italic': r.italic = True
        if kind == 'code':
            r.font.name = 'Consolas'; r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
            r.font.size = Pt((size or 11) - 1)
        if size: r.font.size = Pt(size)
        if color: r.font.color.rgb = color
    return par

def add_field(par, instr, placeholder='', size=None, color=None):
    r = par.add_run()
    fc1 = OxmlElement('w:fldChar'); fc1.set(qn('w:fldCharType'), 'begin'); r._r.append(fc1)
    r2 = par.add_run(); it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = instr; r2._r.append(it)
    r3 = par.add_run(); fc2 = OxmlElement('w:fldChar'); fc2.set(qn('w:fldCharType'), 'separate'); r3._r.append(fc2)
    r4 = par.add_run(placeholder)
    r5 = par.add_run(); fc3 = OxmlElement('w:fldChar'); fc3.set(qn('w:fldCharType'), 'end'); r5._r.append(fc3)
    # Word toma el formato del resultado desde las runs del campo: se aplica a todas
    for rr in (r, r2, r3, r4, r5):
        if size: rr.font.size = Pt(size)
        if color: rr.font.color.rgb = color
    return r4

def nueva_lista_numerada(doc):
    """Crea una instancia de numeración nueva (reinicia en 1) basada en el estilo List Number."""
    numbering = doc.part.numbering_part.numbering_definitions._numbering
    style = doc.styles['List Number']
    num_id_style = style.element.pPr.numPr.numId.val
    abstract_id = None
    for num in numbering.findall(qn('w:num')):
        if int(num.get(qn('w:numId'))) == int(num_id_style):
            abstract_id = int(num.find(qn('w:abstractNumId')).get(qn('w:val'))); break
    num = numbering.add_num(abstract_id)
    num.add_lvlOverride(ilvl=0).add_startOverride(1)
    return num.numId

def set_numbering(par, num_id, ilvl=0):
    pPr = par._p.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    il = OxmlElement('w:ilvl'); il.set(qn('w:val'), str(ilvl)); numPr.append(il)
    ni = OxmlElement('w:numId'); ni.set(qn('w:val'), str(num_id)); numPr.append(ni)
    pPr.append(numPr)

def configurar_estilos(doc):
    st = doc.styles['Normal']; st.font.name = 'Calibri'; st.font.size = Pt(11)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    st.paragraph_format.space_after = Pt(6); st.paragraph_format.line_spacing = 1.12
    for name, size, color, before, after in (('Heading 1', 20, AZUL, 0, 10), ('Heading 2', 14, VERDE, 14, 6), ('Heading 3', 12, GRIS, 10, 4)):
        h = doc.styles[name]; h.font.name = 'Calibri'; h.font.size = Pt(size); h.font.bold = True; h.font.color.rgb = color
        h.element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
        h.paragraph_format.space_before = Pt(before); h.paragraph_format.space_after = Pt(after); h.paragraph_format.keep_with_next = True
    doc.styles['Heading 1'].paragraph_format.page_break_before = True
    for name in ('List Bullet', 'List Number'):
        doc.styles[name].font.name = 'Calibri'; doc.styles[name].font.size = Pt(11); doc.styles[name].paragraph_format.space_after = Pt(3)

def portada(doc, meta):
    sec = doc.sections[0]
    # logos en una tabla sin bordes
    t = doc.add_table(rows=1, cols=2); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c0, c1 = t.rows[0].cells
    c0.paragraphs[0].add_run().add_picture(str(FIG / 'logo-bien-publico.png'), height=Cm(2.6))
    p1 = c1.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p1.add_run().add_picture(str(FIG / 'logo-corfo.png'), height=Cm(2.0)); p1.add_run('    ')
    p1.add_run().add_picture(str(FIG / 'logo-dps.png'), height=Cm(2.0))
    for _ in range(4): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run('MANUAL DE USUARIO'); r.font.size = Pt(12); r.font.color.rgb = VERDE; r.bold = True
    p = doc.add_paragraph(); r = p.add_run(meta.get('titulo', '')); r.font.size = Pt(28); r.bold = True; r.font.color.rgb = AZUL
    p.paragraph_format.space_after = Pt(4)
    p = doc.add_paragraph(); r = p.add_run(meta.get('subtitulo', '')); r.font.size = Pt(14); r.font.color.rgb = GRIS
    # línea
    p = doc.add_paragraph(); pPr = p._p.get_or_add_pPr(); pb = OxmlElement('w:pBdr'); bt = OxmlElement('w:bottom')
    bt.set(qn('w:val'), 'single'); bt.set(qn('w:sz'), '12'); bt.set(qn('w:color'), '0D7377'); bt.set(qn('w:space'), '1'); pb.append(bt); pPr.append(pb)
    for _ in range(2): doc.add_paragraph()
    datos = [('Versión', meta.get('version', '')), ('Fecha', meta.get('fecha', '')), ('Destinatario', meta.get('destinatario', '')),
             ('Elaborado por', meta.get('elaborado', '')), ('Sitio web', meta.get('sitio', ''))]
    tb = doc.add_table(rows=len(datos), cols=2); tb.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (k, v) in enumerate(datos):
        a, b = tb.rows[i].cells; a.width = Cm(3.8); b.width = Cm(12)
        ra = a.paragraphs[0].add_run(k); ra.bold = True; ra.font.size = Pt(10.5); ra.font.color.rgb = AZUL
        rb = b.paragraphs[0].add_run(v); rb.font.size = Pt(10.5)
        for c in (a, b): set_cell_borders(c, color='FFFFFF', sz=0)
    for _ in range(2): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run('Proyecto apoyado por CORFO · Programa Desarrollo Productivo Sostenible'); r.font.size = Pt(9); r.font.color.rgb = GRIS
    doc.add_page_break()

def indice(doc):
    h = doc.add_paragraph(); r = h.add_run('Índice'); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = AZUL
    p = doc.add_paragraph()
    add_field(p, ' TOC \\o "1-3" \\h \\z \\u ', 'Índice: si Word no lo muestra, haga clic derecho aquí y elija "Actualizar campos".')
    # pedir a Word que actualice los campos al abrir
    import os
    if os.environ.get('MANUAL_UPDATEFIELDS') == '1':
        settings = doc.settings.element
        uf = OxmlElement('w:updateFields'); uf.set(qn('w:val'), 'true'); settings.append(uf)
    doc.add_page_break()

def encabezado_pie(doc, meta):
    sec = doc.sections[0]
    hp = sec.header.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run(meta.get('titulo', '')); r.font.size = Pt(8.5); r.font.color.rgb = GRIS
    fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run(f"Versión {meta.get('version','')} · {meta.get('fecha','')} · Página "); r.font.size = Pt(8.5); r.font.color.rgb = GRIS
    add_field(fp, ' PAGE ', '1', size=8.5, color=GRIS)
    r = fp.add_run(' de '); r.font.size = Pt(8.5); r.font.color.rgb = GRIS
    add_field(fp, ' NUMPAGES ', '1', size=8.5, color=GRIS)
    sec.different_first_page_header_footer = True

def figura(doc, b, n):
    path = FIG / b['file']
    if not path.exists():
        print('  ⚠ figura ausente:', b['file']); return
    w, h = Image.open(path).size
    max_w, max_h = Cm(16.0), Cm(19.5)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    if h / w * max_w > max_h:
        p.add_run().add_picture(str(path), height=max_h)
    else:
        p.add_run().add_picture(str(path), width=max_w)
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER; c.paragraph_format.space_after = Pt(10)
    r = c.add_run(f'Figura {n}. '); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = AZUL
    add_runs(c, b['caption'], size=9, color=GRIS)

def tabla(doc, b):
    ncol = len(b['header'])
    t = doc.add_table(rows=1 + len(b['rows']), cols=ncol); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, htxt in enumerate(b['header']):
        cell = t.rows[0].cells[j]; set_cell_shading(cell, '1B3A5C')
        par = cell.paragraphs[0]; add_runs(par, htxt, size=9.5, color=RGBColor(0xFF, 0xFF, 0xFF), bold_all=True)
    for i, row in enumerate(b['rows'], start=1):
        for j, ctxt in enumerate(row):
            cell = t.rows[i].cells[j]
            if i % 2 == 0: set_cell_shading(cell, 'F3F5F7')
            add_runs(cell.paragraphs[0], ctxt, size=9.5)
    # anchos: primera columna más angosta en tablas de 2 columnas
    total = Cm(16.0)
    if ncol == 2:
        widths = [Cm(6.0), Cm(10.0)]
    elif ncol == 3:
        widths = [Cm(4.6), Cm(7.0), Cm(4.4)]
    elif ncol == 4:
        widths = [Cm(2.2), Cm(3.6), Cm(6.2), Cm(4.0)]
    else:
        widths = [total / ncol] * ncol
    for row in t.rows:
        for j, c in enumerate(row.cells): c.width = widths[j]
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def recuadro(doc, b):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]; cell.width = Cm(16.0)
    set_cell_shading(cell, ADMON_COLOR[b['kind']]); set_cell_borders(cell, color=ADMON_COLOR[b['kind']], sz=4, left_color=ADMON_BORDE[b['kind']], left_sz=28)
    par = cell.paragraphs[0]; par.paragraph_format.space_before = Pt(3); par.paragraph_format.space_after = Pt(3)
    r = par.add_run(ADMON_LABEL[b['kind']] + ': '); r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = RGBColor.from_string(ADMON_BORDE[b['kind']])
    add_runs(par, b['text'], size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def main():
    meta, blocks = parse(BASE / 'fuente' / 'manual.md')
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.59), Cm(27.94)  # carta
    sec.left_margin = sec.right_margin = Cm(2.5); sec.top_margin = Cm(2.3); sec.bottom_margin = Cm(2.2)
    configurar_estilos(doc)
    encabezado_pie(doc, meta)
    portada(doc, meta)
    indice(doc)
    nfig = 0
    first_heading = True
    for b in blocks:
        t = b['type']
        if t == 'heading':
            if b['level'] == 1:
                if first_heading:
                    # la primera parte no necesita salto extra: el índice ya terminó en salto de página
                    h = doc.add_heading(b['text'], level=1); h.paragraph_format.page_break_before = False; first_heading = False
                else:
                    doc.add_heading(b['text'], level=1)
            else:
                doc.add_heading(b['text'], level=b['level'])
        elif t == 'para':
            add_runs(doc.add_paragraph(), b['text'])
        elif t == 'list':
            if b['ordered']:
                num_id = nueva_lista_numerada(doc)
                for it in b['items']:
                    p = doc.add_paragraph(style='List Number'); set_numbering(p, num_id); add_runs(p, it)
            else:
                for it in b['items']:
                    add_runs(doc.add_paragraph(style='List Bullet'), it)
        elif t == 'table':
            tabla(doc, b)
        elif t == 'figure':
            nfig += 1; figura(doc, b, nfig)
        elif t == 'admon':
            recuadro(doc, b)
        elif t == 'pagebreak':
            pass  # cada capítulo (Heading 1) ya inicia en página nueva
    doc.core_properties.title = meta.get('titulo', ''); doc.core_properties.subject = meta.get('subtitulo', '')
    doc.core_properties.author = 'Bien Público H2V Araucanía (24BP-269085)'; doc.core_properties.keywords = 'H2V, Araucanía, manual, administración, Payload'
    doc.save(SALIDA)
    print(f'✓ DOCX: {SALIDA} ({SALIDA.stat().st_size/1e6:.1f} MB, {nfig} figuras)')

if __name__ == '__main__':
    main()
